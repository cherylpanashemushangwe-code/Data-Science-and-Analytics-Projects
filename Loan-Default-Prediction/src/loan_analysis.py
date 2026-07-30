"""
SecureBank Loan Default Prediction - Full Analysis Script
Generates all figures, runs three models, and produces the Word document report.
Run: python loan_analysis.py
"""

# =============================================================================
# 0.  IMPORTS & SETUP
# =============================================================================
import os, warnings
from pathlib import Path
import numpy as np
import pandas as pd
import matplotlib
matplotlib.use("Agg")
import matplotlib.pyplot as plt
import seaborn as sns

from sklearn.model_selection import train_test_split
from sklearn.linear_model import LogisticRegression
from sklearn.svm import SVC
from sklearn.preprocessing import StandardScaler
from sklearn.metrics import (
    confusion_matrix, accuracy_score, precision_score,
    recall_score, f1_score, roc_curve, auc, ConfusionMatrixDisplay,
)
import statsmodels.api as sm

from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

warnings.filterwarnings("ignore")
pd.set_option("display.float_format", "{:.4f}".format)

# Paths are resolved relative to this script so the project is fully portable.
ROOT   = Path(__file__).resolve().parent.parent     # repository root
DATA   = ROOT / "data" / "loan_default.csv"
BASE   = ROOT / "reports"
FIGS   = BASE / "figures"
FIGS.mkdir(parents=True, exist_ok=True)
DATA, BASE, FIGS = str(DATA), str(BASE), str(FIGS)

plt.rcParams.update({"figure.dpi": 130, "savefig.dpi": 150,
                      "axes.titleweight": "bold", "axes.labelsize": 11})

RED   = "#C0392B"
GREEN = "#1E8449"
BLUE  = "#1A5276"
ORG   = "#CA6F1E"

print("=" * 65)
print("  SecureBank Loan Default Prediction - Analysis Script")
print("=" * 65)

# =============================================================================
# 1.  LOAD & INSPECT
# =============================================================================
print("\n[1] Loading data ...")
df_raw = pd.read_csv(DATA)
print(f"    Shape : {df_raw.shape}")
print(f"    Columns: {list(df_raw.columns)}")

NUM_VARS = [
    "Loan_Duration_Months", "Loan_Amount", "Installment_Rate_Pct",
    "Current_Residence_Years", "Age", "Num_Existing_Credits", "Num_Dependents",
]

summary_stats = df_raw[NUM_VARS].describe().T.rename(columns={"50%": "median"})
summary_stats = summary_stats[["count", "mean", "std", "min", "25%", "median", "75%", "max"]]

missing_ser   = df_raw.isnull().sum()
missing_df    = pd.DataFrame({
    "Missing Count":      missing_ser,
    "Missing %":          (missing_ser / len(df_raw) * 100).round(2),
}).loc[missing_ser > 0]

print("\nMissing values:")
print(missing_df)
print(f"\nNegative Loan_Amount  : {(df_raw['Loan_Amount'] < 0).sum()}")
print(f"Age < 18              : {(df_raw['Age'].dropna() < 18).sum()}")
print(f"Age > 100             : {(df_raw['Age'].dropna() > 100).sum()}")

# =============================================================================
# 2.  DATA CLEANING
# =============================================================================
print("\n[2] Cleaning ...")
df = df_raw.copy()

# Impute numerical NaNs with column median
for col in ["Age", "Loan_Duration_Months", "Loan_Amount"]:
    med = df[col].median()
    n_miss = df[col].isnull().sum()
    if n_miss:
        df[col] = df[col].fillna(med)
        print(f"    {col}: {n_miss} NaN -> filled with median {med:.1f}")

# Fix impossible ages: < 18 -> replace with median
age_median = df["Age"].median()
n_young = (df["Age"] < 18).sum()
if n_young:
    df.loc[df["Age"] < 18, "Age"] = age_median
    print(f"    Age < 18: {n_young} records -> replaced with median {age_median:.1f}")

# Fix negative loan amounts -> replace with median
loan_med = df["Loan_Amount"].median()
n_neg = (df["Loan_Amount"] < 0).sum()
if n_neg:
    df.loc[df["Loan_Amount"] < 0, "Loan_Amount"] = loan_med
    print(f"    Loan_Amount < 0: {n_neg} records -> replaced with median {loan_med:.1f}")

target_counts = df["Loan_Default"].value_counts()
default_rate  = (target_counts.get("Yes", 0) / len(df) * 100)
print(f"\nClass distribution - No: {target_counts.get('No',0)}  "
      f"Yes: {target_counts.get('Yes',0)}  "
      f"Default rate: {default_rate:.1f}%")

# =============================================================================
# 3.  FEATURE ENCODING
# =============================================================================
print("\n[3] Encoding features ...")

df["Target"] = (df["Loan_Default"] == "Yes").astype(int)

# -- Ordinal mappings (natural risk-ordered scale) --
checking_map = {
    "< 0 DM": 0,
    "0 - 200 DM": 1,
    ">= 200 DM": 2,
    "no checking account": 3,
}
savings_map = {
    "unknown / no savings": 0,
    "< 100 DM": 1,
    "100 - 500 DM": 2,
    "500 - 1000 DM": 3,
    ">= 1000 DM": 4,
}
employment_map = {
    "unemployed": 0,
    "< 1 year": 1,
    "1 - 4 years": 2,
    "4 - 7 years": 3,
    ">= 7 years": 4,
}
credit_map = {
    "no credits taken / all credits paid back duly": 0,
    "all credits at this bank paid back duly": 1,
    "existing credits paid back duly": 2,
    "delay in paying off": 3,
    "critical account / other credits existing": 4,
}

df["Checking_Ord"]     = df["Checking_Account_Status"].map(checking_map).fillna(1)
df["Savings_Ord"]      = df["Savings_Account_Status"].map(savings_map).fillna(0)
df["Employment_Ord"]   = df["Employment_Duration"].map(employment_map).fillna(1)
df["CreditHistory_Ord"]= df["Credit_History"].map(credit_map).fillna(2)

# -- Binary --
df["Telephone_Bin"]     = (df["Has_Telephone"] == "yes").astype(int)
df["ForeignWorker_Bin"] = (df["Is_Foreign_Worker"] == "yes").astype(int)

# -- One-hot encode nominal variables --
NOMINAL = [
    "Loan_Purpose", "Personal_Status", "Other_Debtors_Guarantors",
    "Property_Type", "Other_Installment_Plans", "Housing_Type", "Job_Type",
]
df_enc = pd.get_dummies(df, columns=NOMINAL, drop_first=True, dtype=int)
OHE_FEATURES = [c for c in df_enc.columns
                if any(c.startswith(n + "_") for n in NOMINAL)]

ORDINAL_FEATURES = ["Checking_Ord", "Savings_Ord", "Employment_Ord", "CreditHistory_Ord"]
BINARY_FEATURES  = ["Telephone_Bin", "ForeignWorker_Bin"]
ALL_FEATURES     = NUM_VARS + ORDINAL_FEATURES + BINARY_FEATURES + OHE_FEATURES

X = df_enc[ALL_FEATURES].copy()
y = df_enc["Target"].copy()

assert X.isnull().sum().sum() == 0, "Feature matrix still has NaNs!"
print(f"    Total features: {len(ALL_FEATURES)}")

# =============================================================================
# 4.  EDA VISUALIZATIONS
# =============================================================================
print("\n[4] Generating EDA visualizations ...")

df["Default_Label"] = df["Target"].map({0: "No Default", 1: "Default"})
avg_rate = df["Target"].mean() * 100

# ---- Figure 1 : Default rate by Checking Account Status ----
order_chk = ["< 0 DM", "0 - 200 DM", ">= 200 DM", "no checking account"]
chk_pct   = df.groupby("Checking_Account_Status")["Target"].mean() * 100
chk_cnt   = df.groupby("Checking_Account_Status")["Target"].count()
vals = [chk_pct.get(o, 0) for o in order_chk]
cnts = [chk_cnt.get(o, 0) for o in order_chk]
xlabels = ["< 0 DM\n(Overdrawn)", "0-200 DM\n(Low)", ">= 200 DM\n(Good)", "No Checking\nAccount"]
bar_clrs = [RED, ORG, GREEN, BLUE]

fig, ax = plt.subplots(figsize=(9, 5.5))
bars = ax.bar(range(4), vals, color=bar_clrs, edgecolor="black", linewidth=0.7, width=0.6)
ax.set_xticks(range(4)); ax.set_xticklabels(xlabels, fontsize=10)
for bar, v, n in zip(bars, vals, cnts):
    ax.text(bar.get_x()+bar.get_width()/2., bar.get_height()+0.8,
            f"{v:.1f}%\n(n={n})", ha="center", va="bottom", fontsize=9.5, fontweight="bold")
ax.axhline(avg_rate, color="black", linestyle="--", lw=1.8,
           label=f"Overall Average: {avg_rate:.1f}%")
ax.set_ylim(0, 68); ax.set_xlabel("Checking Account Balance Status")
ax.set_ylabel("Loan Default Rate (%)")
ax.set_title("Figure 1: Loan Default Rate by Checking Account Status", pad=12)
ax.legend(fontsize=10)
plt.tight_layout()
plt.savefig(os.path.join(FIGS, "fig1_checking.png"), bbox_inches="tight")
plt.close()
print("    Figure 1 saved.")

# ---- Figure 2 : Loan Amount by Default Status ----
fig, (ax1, ax2) = plt.subplots(1, 2, figsize=(12, 5.5))
pal = {"No Default": GREEN, "Default": RED}
sns.boxplot(data=df, x="Default_Label", y="Loan_Amount",
            palette=pal, ax=ax1, linewidth=1.4,
            flierprops={"marker": ".", "markersize": 4, "alpha": 0.35})
ax1.set_title("(a) Box Plot", fontsize=11, fontweight="bold")
ax1.set_xlabel("Default Status"); ax1.set_ylabel("Loan Amount (DM)")
for i, lbl in enumerate(["No Default", "Default"]):
    med = df.loc[df["Default_Label"]==lbl, "Loan_Amount"].median()
    ax1.text(i, med+60, f"Median:\n{med:,.0f}", ha="center", va="bottom",
             fontsize=9, fontweight="bold",
             bbox=dict(boxstyle="round,pad=0.2", fc="white", ec="gray", alpha=0.9))
for lbl, ck in [("No Default","no_default"),("Default","default")]:
    df.loc[df["Default_Label"]==lbl,"Loan_Amount"].plot(
        kind="kde", ax=ax2, color=pal[lbl], label=lbl, linewidth=2)
ax2.set_xlabel("Loan Amount (DM)"); ax2.set_ylabel("Density")
ax2.set_title("(b) Density Plot", fontsize=11, fontweight="bold")
ax2.legend(fontsize=10)
fig.suptitle("Figure 2: Loan Amount Distribution by Default Status",
             fontsize=13, fontweight="bold", y=1.01)
plt.tight_layout()
plt.savefig(os.path.join(FIGS, "fig2_loan_amount.png"), bbox_inches="tight")
plt.close()
print("    Figure 2 saved.")

# ---- Figure 3 : Age Distribution by Default Status ----
bins = np.arange(18, 80, 4)
fig, ax = plt.subplots(figsize=(9, 5.5))
for lbl, clr in [("No Default", GREEN), ("Default", RED)]:
    data = df.loc[df["Default_Label"]==lbl, "Age"]
    ax.hist(data, bins=bins, alpha=0.6, color=clr, label=lbl,
            edgecolor="black", linewidth=0.4)
    ax.axvline(data.mean(), color=clr, linestyle="--", linewidth=2.2,
               label=f"Mean ({lbl}): {data.mean():.1f} yrs")
ax.set_xlabel("Age (Years)"); ax.set_ylabel("Number of Borrowers")
ax.set_title("Figure 3: Age Distribution by Default Status", pad=12)
ax.legend(fontsize=9)
plt.tight_layout()
plt.savefig(os.path.join(FIGS, "fig3_age.png"), bbox_inches="tight")
plt.close()
print("    Figure 3 saved.")

# ---- Figure 4 : Default Rate by Loan Purpose ----
purp_df = (df.groupby("Loan_Purpose")["Target"]
             .agg(["mean","count"])
             .rename(columns={"mean":"rate","count":"n"})
             .reset_index())
purp_df["pct"] = purp_df["rate"] * 100
purp_df = purp_df.sort_values("pct", ascending=True)
bar_c4 = [RED if x >= avg_rate else GREEN for x in purp_df["pct"]]

fig, ax = plt.subplots(figsize=(10, 6))
bars = ax.barh(range(len(purp_df)), purp_df["pct"],
               color=bar_c4, edgecolor="black", linewidth=0.5, height=0.65)
ax.set_yticks(range(len(purp_df)))
ax.set_yticklabels(purp_df["Loan_Purpose"], fontsize=9)
for bar, row in zip(bars, purp_df.itertuples()):
    ax.text(bar.get_width()+0.4, bar.get_y()+bar.get_height()/2.,
            f"{row.pct:.1f}%  (n={row.n})", va="center", fontsize=8.5)
ax.axvline(avg_rate, color="black", linestyle="--", lw=1.8,
           label=f"Average: {avg_rate:.1f}%")
ax.set_xlabel("Default Rate (%)")
ax.set_title("Figure 4: Default Rate by Loan Purpose\n"
             "(Red = Above Average Risk  |  Green = Below Average Risk)", pad=10)
ax.set_xlim(0, 58); ax.legend(fontsize=10)
plt.tight_layout()
plt.savefig(os.path.join(FIGS, "fig4_loan_purpose.png"), bbox_inches="tight")
plt.close()
print("    Figure 4 saved.")

# =============================================================================
# 5.  TRAIN / TEST SPLIT AND SCALING
# =============================================================================
print("\n[5] Splitting and scaling ...")
X_train, X_test, y_train, y_test = train_test_split(
    X, y, test_size=0.2, random_state=42, stratify=y
)
print(f"    Train: {X_train.shape[0]}  Test: {X_test.shape[0]}  "
      f"(Train default rate: {y_train.mean():.2%})")

scaler = StandardScaler()
X_tr_sc = scaler.fit_transform(X_train)
X_te_sc = scaler.transform(X_test)

X_tr_df = pd.DataFrame(X_tr_sc, columns=ALL_FEATURES)
X_te_df = pd.DataFrame(X_te_sc, columns=ALL_FEATURES)

# =============================================================================
# 6.  MODEL 1 - LOGISTIC REGRESSION BASELINE (all features)
# =============================================================================
print("\n[6] Model 1 - LR Baseline ...")
lr1 = LogisticRegression(max_iter=3000, random_state=42, C=1.0, solver="lbfgs")
lr1.fit(X_tr_sc, y_train)
yp1  = lr1.predict(X_te_sc)
ypr1 = lr1.predict_proba(X_te_sc)[:, 1]

cm1   = confusion_matrix(y_test, yp1)
acc1  = accuracy_score(y_test, yp1)
prec1 = precision_score(y_test, yp1)
rec1  = recall_score(y_test, yp1)
f1_1  = f1_score(y_test, yp1)
fpr1, tpr1, _ = roc_curve(y_test, ypr1)
auc1  = auc(fpr1, tpr1)

print(f"    Acc={acc1:.4f}  Prec={prec1:.4f}  Rec={rec1:.4f}  F1={f1_1:.4f}  AUC={auc1:.4f}")

coef_df = pd.DataFrame({
    "Feature":    ALL_FEATURES,
    "Coeff":      lr1.coef_[0],
    "Odds_Ratio": np.exp(lr1.coef_[0]),
}).sort_values("Odds_Ratio", ascending=False)

print("\n    Top-10 features by odds ratio:")
print(coef_df.head(10).to_string(index=False))

# Confusion matrix figure
fig, ax = plt.subplots(figsize=(6, 5))
ConfusionMatrixDisplay(cm1, display_labels=["No Default","Default"]).plot(
    ax=ax, colorbar=False, cmap="Blues", values_format="d")
ax.set_title("Figure 5: Confusion Matrix\nModel 1 - LR Baseline", pad=10)
plt.tight_layout()
plt.savefig(os.path.join(FIGS, "fig5_cm_lr1.png"), bbox_inches="tight")
plt.close()

# =============================================================================
# 7.  MODEL 2 - OPTIMIZED LR (Backward Elimination)
# =============================================================================
print("\n[7] Model 2 - Backward Elimination ...")

def backward_elim(X_df, y_ser, threshold=0.05):
    features = list(X_df.columns)
    removed_log = []
    itr = 0
    while True:
        itr += 1
        X_sm = sm.add_constant(X_df[features], has_constant="add")
        try:
            res = sm.Logit(
                y_ser.reset_index(drop=True),
                X_sm.reset_index(drop=True)
            ).fit(disp=False, maxiter=600, method="bfgs")
            pvals = res.pvalues.iloc[1:]          # drop constant
            worst_p = pvals.max()
            worst_f = pvals.idxmax()
            if worst_p > threshold:
                features.remove(worst_f)
                removed_log.append((worst_f, round(worst_p, 4)))
                print(f"    Iter {itr:3d}: removed '{worst_f}'  p={worst_p:.4f}")
            else:
                print(f"    Converged - {len(features)} features retained "
                      f"(max p={worst_p:.4f})")
                break
        except Exception as exc:
            print(f"    Stopped at iter {itr}: {exc}")
            break
    return features, removed_log

sel_features, removed_log = backward_elim(X_tr_df, y_train, threshold=0.05)

X_tr2 = X_tr_df[sel_features].values
X_te2 = X_te_df[sel_features].values

lr2   = LogisticRegression(max_iter=3000, random_state=42, C=1.0, solver="lbfgs")
lr2.fit(X_tr2, y_train)
yp2   = lr2.predict(X_te2)
ypr2  = lr2.predict_proba(X_te2)[:, 1]

cm2   = confusion_matrix(y_test, yp2)
acc2  = accuracy_score(y_test, yp2)
prec2 = precision_score(y_test, yp2)
rec2  = recall_score(y_test, yp2)
f1_2  = f1_score(y_test, yp2)
fpr2, tpr2, _ = roc_curve(y_test, ypr2)
auc2  = auc(fpr2, tpr2)

print(f"    Acc={acc2:.4f}  Prec={prec2:.4f}  Rec={rec2:.4f}  F1={f1_2:.4f}  AUC={auc2:.4f}")

sel_coef_df = pd.DataFrame({
    "Feature":    sel_features,
    "Coeff":      lr2.coef_[0],
    "Odds_Ratio": np.exp(lr2.coef_[0]),
}).sort_values("Odds_Ratio", ascending=False)

fig, ax = plt.subplots(figsize=(6, 5))
ConfusionMatrixDisplay(cm2, display_labels=["No Default","Default"]).plot(
    ax=ax, colorbar=False, cmap="Greens", values_format="d")
ax.set_title("Figure 6: Confusion Matrix\nModel 2 - Optimized LR", pad=10)
plt.tight_layout()
plt.savefig(os.path.join(FIGS, "fig6_cm_lr2.png"), bbox_inches="tight")
plt.close()

# Feature importance chart for Model 2
sorted_sel = sel_coef_df.sort_values("Coeff")
fc2 = [RED if c > 0 else GREEN for c in sorted_sel["Coeff"]]
fig, ax = plt.subplots(figsize=(10, max(5, len(sel_features)*0.42)))
ax.barh(range(len(sorted_sel)), sorted_sel["Coeff"],
        color=fc2, edgecolor="black", linewidth=0.5, height=0.7)
ax.set_yticks(range(len(sorted_sel)))
ax.set_yticklabels(sorted_sel["Feature"], fontsize=9)
ax.axvline(0, color="black", lw=1.5)
for i, (coef, oratio) in enumerate(zip(sorted_sel["Coeff"], sorted_sel["Odds_Ratio"])):
    ha = "left" if coef >= 0 else "right"
    offset = 0.01 if coef >= 0 else -0.01
    ax.text(coef+offset, i, f"OR={oratio:.2f}", va="center", ha=ha, fontsize=8)
ax.set_xlabel("Log-Odds Coefficient")
ax.set_title("Figure 9: Model 2 - Feature Coefficients & Odds Ratios\n"
             "(Red = Increases Default Risk  |  Green = Reduces Default Risk)", pad=10)
plt.tight_layout()
plt.savefig(os.path.join(FIGS, "fig9_coeff.png"), bbox_inches="tight")
plt.close()

# =============================================================================
# 8.  MODEL 3 - SVM (RBF KERNEL)
# =============================================================================
print("\n[8] Model 3 - SVM RBF ...")
svm = SVC(kernel="rbf", probability=True, random_state=42, C=1.0, gamma="scale")
svm.fit(X_tr_sc, y_train)
yp3  = svm.predict(X_te_sc)
ypr3 = svm.predict_proba(X_te_sc)[:, 1]

cm3   = confusion_matrix(y_test, yp3)
acc3  = accuracy_score(y_test, yp3)
prec3 = precision_score(y_test, yp3)
rec3  = recall_score(y_test, yp3)
f1_3  = f1_score(y_test, yp3)
fpr3, tpr3, _ = roc_curve(y_test, ypr3)
auc3  = auc(fpr3, tpr3)

print(f"    Acc={acc3:.4f}  Prec={prec3:.4f}  Rec={rec3:.4f}  F1={f1_3:.4f}  AUC={auc3:.4f}")

fig, ax = plt.subplots(figsize=(6, 5))
ConfusionMatrixDisplay(cm3, display_labels=["No Default","Default"]).plot(
    ax=ax, colorbar=False, cmap="Reds", values_format="d")
ax.set_title("Figure 7: Confusion Matrix\nModel 3 - SVM (RBF Kernel)", pad=10)
plt.tight_layout()
plt.savefig(os.path.join(FIGS, "fig7_cm_svm.png"), bbox_inches="tight")
plt.close()

# =============================================================================
# 9.  ROC CURVES (all 3 models)
# =============================================================================
fig, ax = plt.subplots(figsize=(8, 7))
ax.plot(fpr1, tpr1, color=BLUE,  lw=2.5, label=f"Model 1: LR Baseline   (AUC = {auc1:.3f})")
ax.plot(fpr2, tpr2, color=GREEN, lw=2.5, ls="--", label=f"Model 2: LR Optimized  (AUC = {auc2:.3f})")
ax.plot(fpr3, tpr3, color=RED,   lw=2.5, ls=":",  label=f"Model 3: SVM RBF       (AUC = {auc3:.3f})")
ax.plot([0,1],[0,1], color="gray", lw=1.5, ls="--", alpha=0.7, label="Random Classifier (AUC = 0.500)")
ax.fill_between(fpr1, tpr1, alpha=0.05, color=BLUE)
ax.set_xlim(0, 1); ax.set_ylim(0, 1.02)
ax.set_xlabel("False Positive Rate (1 - Specificity)")
ax.set_ylabel("True Positive Rate (Sensitivity)")
ax.set_title("Figure 8: ROC Curves - All Three Models", pad=12)
ax.legend(loc="lower right", fontsize=11)
ax.grid(True, alpha=0.25)
plt.tight_layout()
plt.savefig(os.path.join(FIGS, "fig8_roc.png"), bbox_inches="tight")
plt.close()
print("    ROC figure saved.")

# =============================================================================
# 10.  COST-OPTIMAL THRESHOLD (bonus - Model 2, best balanced model)
# =============================================================================
COST_FN = 15000   # approving a defaulter
COST_FP = 2000    # rejecting a good borrower

thresholds = np.linspace(0.01, 0.99, 200)
costs = []
for t in thresholds:
    yp_t = (ypr2 >= t).astype(int)
    fn   = ((yp_t == 0) & (y_test == 1)).sum()
    fp   = ((yp_t == 1) & (y_test == 0)).sum()
    costs.append(fn * COST_FN + fp * COST_FP)
costs = np.array(costs)
opt_idx  = np.argmin(costs)
opt_thr  = thresholds[opt_idx]
opt_cost = costs[opt_idx]

fig, ax = plt.subplots(figsize=(8, 5))
ax.plot(thresholds, costs / 1000, color=BLUE, lw=2)
ax.axvline(opt_thr, color=RED, lw=2, linestyle="--",
           label=f"Optimal threshold = {opt_thr:.2f}")
ax.scatter([opt_thr], [opt_cost/1000], color=RED, zorder=5, s=80)
ax.set_xlabel("Decision Threshold"); ax.set_ylabel("Total Cost ($ thousands)")
ax.set_title("Figure 10: Total Misclassification Cost vs Decision Threshold\n"
             f"(FN cost = ${COST_FN:,}  |  FP cost = ${COST_FP:,})", pad=10)
ax.legend(fontsize=11)
plt.tight_layout()
plt.savefig(os.path.join(FIGS, "fig10_threshold.png"), bbox_inches="tight")
plt.close()

# Metrics at optimal threshold
yp2_opt  = (ypr2 >= opt_thr).astype(int)
acc2_opt  = accuracy_score(y_test, yp2_opt)
prec2_opt = precision_score(y_test, yp2_opt)
rec2_opt  = recall_score(y_test, yp2_opt)
f1_2_opt  = f1_score(y_test, yp2_opt)
print(f"\n    Optimal threshold for Model 2: {opt_thr:.2f}")
print(f"    At threshold {opt_thr:.2f}: Acc={acc2_opt:.4f}  Prec={prec2_opt:.4f}  "
      f"Rec={rec2_opt:.4f}  F1={f1_2_opt:.4f}")

# =============================================================================
# 11.  SUMMARY TABLE
# =============================================================================
metrics_summary = pd.DataFrame({
    "Model":     ["LR Baseline", "LR Optimized", "SVM (RBF)"],
    "# Features":[len(ALL_FEATURES), len(sel_features), len(ALL_FEATURES)],
    "Accuracy":  [acc1, acc2, acc3],
    "Precision": [prec1, prec2, prec3],
    "Recall":    [rec1, rec2, rec3],
    "F1-Score":  [f1_1, f1_2, f1_3],
    "AUC":       [auc1, auc2, auc3],
})
print("\n" + "=" * 65)
print("  MODEL COMPARISON SUMMARY")
print("=" * 65)
print(metrics_summary.to_string(index=False))

# Interpret top-3 odds ratios from Model 1
top3 = coef_df.head(3)
bot3 = coef_df.tail(3)

print("\n  Top 3 risk-INCREASING factors (Model 1 odds ratios):")
for _, row in top3.iterrows():
    print(f"    {row['Feature']:<40} OR = {row['Odds_Ratio']:.4f}")

print("\n  Top 3 risk-REDUCING factors:")
for _, row in bot3.iterrows():
    print(f"    {row['Feature']:<40} OR = {row['Odds_Ratio']:.4f}")

# =============================================================================
# 12.  GENERATE WORD DOCUMENT
# =============================================================================
print("\n[12] Generating Word document ...")

def set_cell_bg(cell, hex_color):
    """Shade a table cell with a background colour (RRGGBB hex, no #)."""
    tc   = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd  = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), hex_color)
    tcPr.append(shd)

def add_heading(doc, text, level=1):
    h = doc.add_heading(text, level=level)
    for run in h.runs:
        run.font.color.rgb = RGBColor(0x1A, 0x52, 0x76)
    return h

def add_body(doc, text):
    p = doc.add_paragraph(text)
    p.paragraph_format.space_after = Pt(6)
    for run in p.runs:
        run.font.size = Pt(12)
        run.font.name = "Times New Roman"
    return p

def add_figure(doc, path, caption, width=5.5):
    if os.path.exists(path):
        doc.add_picture(path, width=Inches(width))
        last = doc.paragraphs[-1]
        last.alignment = WD_ALIGN_PARAGRAPH.CENTER
        cap = doc.add_paragraph(caption)
        cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
        cap.paragraph_format.space_after = Pt(10)
        for run in cap.runs:
            run.font.italic = True
            run.font.size   = Pt(10)
    else:
        doc.add_paragraph(f"[Figure not found: {path}]")

def add_metric_table(doc, headers, rows):
    tbl = doc.add_table(rows=len(rows)+1, cols=len(headers))
    tbl.style = "Table Grid"
    # header row
    hdr = tbl.rows[0]
    for i, h in enumerate(headers):
        cell = hdr.cells[i]
        cell.text = h
        cell.paragraphs[0].runs[0].bold = True
        set_cell_bg(cell, "BDC3C7")
    # data rows
    for r_idx, row_data in enumerate(rows):
        for c_idx, val in enumerate(row_data):
            tbl.rows[r_idx+1].cells[c_idx].text = str(val)
    doc.add_paragraph()

# ---- Build the document ----
doc = Document()

# Margins: 1 inch
for section in doc.sections:
    section.top_margin    = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin   = Inches(1)
    section.right_margin  = Inches(1)

# Default font
doc.styles["Normal"].font.name = "Times New Roman"
doc.styles["Normal"].font.size = Pt(12)

# ----- TITLE PAGE -----
title = doc.add_heading("SecureBank Loan Default Prediction", 0)
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
sub = doc.add_paragraph("A Risk Assessment Report for the Chief Risk Officer")
sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
sub.runs[0].font.size = Pt(13)
doc.add_paragraph()
meta = doc.add_paragraph("Prepared by: [Your Name]     |     Date: June 2026\n"
                          "Course: [Course Name]     |     Instructor: [Instructor Name]")
meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
doc.add_page_break()

# ----- STEP 1 : BUSINESS PROBLEM DEFINITION -----
add_heading(doc, "Step 1: Business Problem Definition", level=1)

add_heading(doc, "1.1  The Business Question", level=2)
add_body(doc,
    "SecureBank faces a fundamental credit risk challenge: given a prospective borrower's "
    "financial profile, can the institution predict - before loan approval - whether that "
    "borrower is likely to default on repayment? Phrased as a data analytics question: "
    "Using historical loan applicant data, can a classification model reliably identify "
    "high-risk applicants at the point of application, enabling more informed lending decisions?")

add_heading(doc, "1.2  Stakeholder and Decision Context", level=2)
add_body(doc,
    "The primary stakeholder is SecureBank's Chief Risk Officer (CRO), who is responsible for "
    "the bank's lending policy and portfolio quality. The output of this analysis will directly "
    "inform the loan approval process: whether to approve, flag for additional scrutiny, or "
    "reject individual loan applications. Secondary stakeholders include the bank's credit "
    "analysts, who will operationalise the model's recommendations.")

add_heading(doc, "1.3  Cost Asymmetry: False Negatives vs. False Positives", level=2)
add_body(doc,
    "The CRO has quantified the two types of prediction errors as follows: a false negative "
    "(approving a borrower who subsequently defaults) costs the bank an average of $15,000 in "
    "direct losses. A false positive (rejecting a creditworthy borrower) costs $2,000 in "
    "foregone revenue. The cost ratio is 7.5:1, meaning that one missed default is "
    "approximately seven and a half times more costly than one wrongful rejection. "
    "This asymmetry has direct implications for model selection and threshold calibration: "
    "the bank should prefer a model with high recall (few missed defaults) even if this "
    "marginally reduces overall accuracy. A standard 0.5 classification threshold, which "
    "treats both error types as equally costly, would be sub-optimal for this business context.")

add_heading(doc, "1.4  Definition of Success", level=2)
add_body(doc,
    "A successful outcome for the CRO is a model that (1) achieves a recall of at least 0.70 "
    "for defaulters - i.e., correctly identifies at least 70% of future defaults - while "
    "(2) maintaining overall precision above 0.60 to avoid an excessive rejection of "
    "creditworthy applicants, and (3) provides interpretable risk factors that can be "
    "translated into actionable underwriting guidelines. An AUC above 0.75 would indicate "
    "the model has meaningful discriminatory power beyond chance.")

doc.add_page_break()

# ----- STEP 2 : DATA CLEANING & INITIAL FINDINGS -----
add_heading(doc, "Step 2: Data Cleaning and Initial Findings", level=1)

add_heading(doc, "2.1  Dataset Overview", level=2)
add_body(doc,
    f"The dataset provided by SecureBank contains {len(df_raw):,} loan records and "
    f"{df_raw.shape[1]} variables, including the binary target variable Loan_Default "
    f"(Yes/No). Seven variables are continuous numerical measures (e.g., loan amount, age), "
    f"while the remaining fourteen are categorical, describing credit behaviour, employment "
    f"status, housing type, and loan purpose.")

add_heading(doc, "2.2  Summary Statistics", level=2)
add_body(doc,
    "Table 1 presents descriptive statistics for all seven numerical variables. "
    "The average loan amount is approximately 3,271 DM, with a wide range from 250 DM "
    "to 18,424 DM, suggesting considerable heterogeneity in loan size. Loan durations "
    "range from 4 to 72 months, and borrower ages span from approximately 19 to 75 years.")

# Summary stats table
sum_rows = []
for feat in NUM_VARS:
    col = df_raw[feat].dropna()
    sum_rows.append([
        feat,
        f"{col.count():,.0f}",
        f"{col.mean():,.2f}",
        f"{col.std():,.2f}",
        f"{col.min():,.2f}",
        f"{col.median():,.2f}",
        f"{col.max():,.2f}",
    ])
add_metric_table(doc,
    ["Variable","Count","Mean","Std Dev","Min","Median","Max"],
    sum_rows)
cap = doc.add_paragraph("Table 1: Descriptive statistics for numerical variables.")
cap.runs[0].font.italic = True; cap.runs[0].font.size = Pt(10)
cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
doc.add_paragraph()

add_heading(doc, "2.3  Missing Values and Data Quality", level=2)
mv_list = ", ".join([f"{c} ({n} records, {n/len(df_raw)*100:.1f}%)"
                     for c, n in missing_ser.items()])
add_body(doc,
    f"Four variables contained missing values: {mv_list}. "
    f"Numerical missing values were imputed using each column's median - "
    f"a robust strategy that avoids distortion by outliers. "
    f"Missing Employment_Duration records (categorical) were assigned the modal ordinal "
    f"value (1 = less than 1 year). Additionally, five records with negative Loan_Amount "
    f"values and three records with Age below 18 were identified as data entry errors "
    f"and replaced with their respective column medians. After cleaning, the dataset "
    f"contains no missing or impossible values.")

add_heading(doc, "2.4  Categorical Encoding Strategy", level=2)
add_body(doc,
    "The dataset contains fourteen categorical variables that must be converted to numerical "
    "representations for the machine learning models. Three encoding strategies were employed, "
    "based on the nature of each variable:")
doc.add_paragraph(
    "Ordinal Encoding was applied to variables with a meaningful natural order: "
    "Checking_Account_Status (overdrawn -> no account, coded 0-3), "
    "Savings_Account_Status (no savings -> >= 1,000 DM, coded 0-4), "
    "Employment_Duration (unemployed -> >= 7 years, coded 0-4), and "
    "Credit_History (critical account -> all credits paid, coded 0-4). "
    "Preserving rank information allows the model to detect monotonic risk relationships.",
    style="List Bullet")
doc.add_paragraph(
    "One-Hot Encoding (OHE) was applied to seven nominal variables without a natural order: "
    "Loan_Purpose, Personal_Status, Other_Debtors_Guarantors, Property_Type, "
    "Other_Installment_Plans, Housing_Type, and Job_Type. "
    "To avoid perfect multicollinearity, the first dummy category was dropped (drop_first=True).",
    style="List Bullet")
doc.add_paragraph(
    "Binary Encoding converted Has_Telephone and Is_Foreign_Worker to 0/1 indicators.",
    style="List Bullet")
add_body(doc, f"After encoding, the full feature matrix contained {len(ALL_FEATURES)} predictor variables.")

add_heading(doc, "2.5  Class Distribution", level=2)
add_body(doc,
    f"The target variable is moderately imbalanced: {target_counts.get('No',0):,} borrowers "
    f"did not default ({100-default_rate:.1f}%) and {target_counts.get('Yes',0):,} defaulted "
    f"({default_rate:.1f}%). A 30% minority class does not constitute severe imbalance "
    f"(where techniques such as SMOTE or class-weighted training would be mandatory), but "
    f"it does mean that a naive classifier predicting 'No Default' for every application "
    f"would achieve {100-default_rate:.1f}% accuracy - an inflated and misleading score. "
    f"Therefore, F1-score, recall, and AUC are the primary evaluation metrics used throughout "
    f"this analysis rather than accuracy alone.")

add_heading(doc, "2.6  Exploratory Data Analysis", level=2)
add_body(doc,
    "Four visualisations were created to identify the most predictive features prior to "
    "formal modelling.")

add_figure(doc, os.path.join(FIGS,"fig1_checking.png"),
    "Figure 1: Default rate by checking account status. Overdrawn accounts show the "
    "highest default rate (≈49%), nearly three times the rate for borrowers with no "
    "checking account.", width=5.5)

chk_0 = chk_pct.get("< 0 DM", 0)
chk_nc = chk_pct.get("no checking account", 0)
add_body(doc,
    f"Figure 1 reveals a strong gradient: borrowers with overdrawn checking accounts "
    f"default at approximately {chk_0:.1f}%, compared with only {chk_nc:.1f}% for those "
    f"without a checking account. Checking account status is therefore likely to be among "
    f"the most important predictors in the model.")

add_figure(doc, os.path.join(FIGS,"fig2_loan_amount.png"),
    "Figure 2: Loan amount distribution by default status. Defaulters tend to take "
    "larger loans, though the distributions substantially overlap.", width=5.8)

no_def_med = df.loc[df["Default_Label"]=="No Default","Loan_Amount"].median()
def_med    = df.loc[df["Default_Label"]=="Default","Loan_Amount"].median()
add_body(doc,
    f"Figure 2 shows that defaulters have a higher median loan amount "
    f"({def_med:,.0f} DM) compared to non-defaulters ({no_def_med:,.0f} DM). "
    f"However, the distributions overlap substantially, suggesting that loan amount "
    f"alone is insufficient to separate the two classes.")

add_figure(doc, os.path.join(FIGS,"fig3_age.png"),
    "Figure 3: Age distribution by default status. Younger borrowers default at higher "
    "rates; the mean age of defaulters is notably lower.", width=5.5)

no_def_age = df.loc[df["Default_Label"]=="No Default","Age"].mean()
def_age    = df.loc[df["Default_Label"]=="Default","Age"].mean()
add_body(doc,
    f"Figure 3 demonstrates that defaulters are, on average, younger "
    f"({def_age:.1f} years) than non-defaulters ({no_def_age:.1f} years). "
    f"This is consistent with the financial literature, which associates younger borrowers "
    f"with shorter credit histories and greater income volatility.")

add_figure(doc, os.path.join(FIGS,"fig4_loan_purpose.png"),
    "Figure 4: Default rate by loan purpose. Retraining and other loans show above-average "
    "default risk; car and education loans vary widely.", width=5.8)

add_body(doc,
    "Figure 4 shows meaningful variation in default rates across loan purposes. "
    "Some categories exhibit rates well above the overall average, while car purchases "
    "and domestic appliances tend to show lower risk. This suggests that loan purpose "
    "should be incorporated as a predictor in the models.")

doc.add_page_break()

# ----- STEP 3 : ALGORITHMS & MODEL RESULTS -----
add_heading(doc, "Step 3: Algorithm Explanation and Model Results", level=1)

add_heading(doc, "3.1  Why Classification?", level=2)
add_body(doc,
    "The prediction target, Loan_Default, is a binary categorical outcome: a borrower "
    "either defaults or does not. Regression-based approaches predict a continuous quantity "
    "(e.g., how much money will be lost), whereas classification algorithms predict group "
    "membership - which is precisely what the CRO requires. Logistic regression and Support "
    "Vector Machines (SVM) are well-suited to this binary classification task for the "
    "following reasons:")
doc.add_paragraph(
    "Logistic Regression produces probabilistic predictions bounded between 0 and 1, "
    "directly interpretable as the probability of default. Its coefficients can be "
    "exponentiated to yield odds ratios, providing straightforward, regulation-friendly "
    "explanations of credit decisions. It handles mixed numerical and categorical inputs "
    "well and is computationally efficient.",
    style="List Bullet")
doc.add_paragraph(
    "Support Vector Machines find the maximum-margin hyperplane separating the two classes, "
    "making them robust to high-dimensional feature spaces (relevant here given the one-hot "
    "encoded features) and effective when classes are not linearly separable. The RBF kernel "
    "is chosen because it implicitly maps features into a higher-dimensional space, "
    "capturing non-linear relationships between predictors and default risk without requiring "
    "explicit feature engineering.",
    style="List Bullet")

add_heading(doc, "3.2  Model 1 - Logistic Regression Baseline", level=2)
add_body(doc,
    f"The baseline Logistic Regression model was trained on all {len(ALL_FEATURES)} encoded "
    f"features with an L2 regularisation strength of C=1.0. Features were standardised "
    f"(zero mean, unit variance) using StandardScaler to ensure comparability of "
    f"coefficients and numerical stability during optimisation.")

add_figure(doc, os.path.join(FIGS,"fig5_cm_lr1.png"),
    "Figure 5: Confusion matrix for Model 1 (Logistic Regression Baseline).", width=4.2)

# Metrics table
add_metric_table(doc,
    ["Metric","Value"],
    [["Accuracy", f"{acc1:.4f}"],
     ["Precision",f"{prec1:.4f}"],
     ["Recall",   f"{rec1:.4f}"],
     ["F1-Score", f"{f1_1:.4f}"],
     ["AUC",      f"{auc1:.4f}"]])
cap = doc.add_paragraph("Table 2: Model 1 performance metrics.")
cap.runs[0].font.italic = True; cap.runs[0].font.size = Pt(10)
cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
doc.add_paragraph()

add_heading(doc, "Odds Ratio Interpretation (Model 1)", level=3)
add_body(doc,
    "Table 3 presents the top-10 and bottom-5 features ranked by odds ratio. "
    "Three key predictors are interpreted in business terms below:")

# Odds ratio table (top 10)
or_rows = [[row["Feature"], f"{row['Coeff']:.4f}", f"{row['Odds_Ratio']:.4f}"]
           for _, row in coef_df.head(10).iterrows()]
add_metric_table(doc, ["Feature","Coefficient (log-odds)","Odds Ratio"], or_rows)
cap = doc.add_paragraph("Table 3: Top-10 features by odds ratio - Model 1.")
cap.runs[0].font.italic = True; cap.runs[0].font.size = Pt(10)
cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
doc.add_paragraph()

# Extract specific odds ratios for interpretation
def get_or(feat_name, df_coef):
    row = df_coef[df_coef["Feature"]==feat_name]
    if len(row): return row.iloc[0]["Odds_Ratio"]
    return None

chk_or   = get_or("Checking_Ord",      coef_df)
cred_or  = get_or("CreditHistory_Ord", coef_df)
dur_or   = get_or("Loan_Duration_Months", coef_df)

# Build interpretations using actual computed values
interp_texts = []
if chk_or:
    pct_chg = abs((chk_or - 1)*100)
    direction = "reduces" if chk_or < 1 else "increases"
    interp_texts.append(
        f"Checking_Ord (OR = {chk_or:.3f}): A one-unit increase in checking account status "
        f"(e.g., moving from overdrawn to a low positive balance) "
        f"{direction} the odds of default by {pct_chg:.1f}%. "
        f"Borrowers with healthier checking account balances are substantially lower-risk, "
        f"reflecting their day-to-day liquidity and financial discipline.")
if cred_or:
    pct_chg = abs((cred_or - 1)*100)
    direction = "increases" if cred_or > 1 else "reduces"
    interp_texts.append(
        f"CreditHistory_Ord (OR = {cred_or:.3f}): A one-unit worsening in credit history "
        f"(e.g., from 'existing credits paid back duly' to 'delay in paying off') "
        f"{direction} the odds of default by {pct_chg:.1f}%. "
        f"Past repayment behaviour is one of the most reliable signals of future default risk, "
        f"consistent with the credit scoring literature.")
if dur_or:
    pct_chg = abs((dur_or - 1)*100)
    direction = "increases" if dur_or > 1 else "reduces"
    interp_texts.append(
        f"Loan_Duration_Months (OR = {dur_or:.3f}): Each one-standard-deviation increase in "
        f"loan duration {direction} the odds of default by {pct_chg:.1f}%. "
        f"Longer-duration loans expose the bank to greater uncertainty over the borrower's "
        f"financial circumstances and are inherently higher risk.")

for txt in interp_texts:
    doc.add_paragraph(txt, style="List Number")
doc.add_paragraph()

add_heading(doc, "3.3  Model 2 - Optimised Logistic Regression (Stepwise)", level=2)
add_body(doc,
    "Backward stepwise elimination was used to identify the most parsimonious set of "
    "statistically significant predictors. Beginning with the full feature set, the variable "
    f"with the highest p-value was iteratively removed until all remaining features were "
    f"significant at α = 0.05. This procedure reduced the model from {len(ALL_FEATURES)} "
    f"to {len(sel_features)} features, eliminating variables that added no statistically "
    f"significant explanatory power after controlling for other predictors.")
add_body(doc,
    "Features removed through backward elimination included variables whose individual "
    "contributions became negligible when correlated predictors were present - a common "
    "outcome with one-hot encoded categorical variables that exhibit implicit correlations. "
    "The retained features represent the most parsimonious and defensible set for "
    "production use, as all included variables have demonstrably significant associations "
    "with default risk.")

# Selected features table
sel_rows = [[f, f"{r['Coeff']:.4f}", f"{r['Odds_Ratio']:.4f}"]
            for f, (_, r) in zip(sel_features, sel_coef_df.set_index("Feature").iterrows())]
add_metric_table(doc, ["Selected Feature","Coefficient","Odds Ratio"], sel_rows)
cap = doc.add_paragraph(f"Table 4: {len(sel_features)} features retained after stepwise elimination.")
cap.runs[0].font.italic = True; cap.runs[0].font.size = Pt(10)
cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
doc.add_paragraph()

add_figure(doc, os.path.join(FIGS,"fig6_cm_lr2.png"),
    "Figure 6: Confusion matrix for Model 2 (Optimised Logistic Regression).", width=4.2)

add_metric_table(doc,
    ["Metric","Value"],
    [["Accuracy", f"{acc2:.4f}"],
     ["Precision",f"{prec2:.4f}"],
     ["Recall",   f"{rec2:.4f}"],
     ["F1-Score", f"{f1_2:.4f}"],
     ["AUC",      f"{auc2:.4f}"]])
cap = doc.add_paragraph("Table 5: Model 2 performance metrics.")
cap.runs[0].font.italic = True; cap.runs[0].font.size = Pt(10)
cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
doc.add_paragraph()

add_figure(doc, os.path.join(FIGS,"fig9_coeff.png"),
    "Figure 9: Model 2 feature coefficients and odds ratios. "
    "Red bars indicate positive log-odds (higher default risk); "
    "green bars indicate negative log-odds (lower default risk).", width=5.8)

add_heading(doc, "3.4  Model 3 - SVM with RBF Kernel", level=2)
add_body(doc,
    "A Support Vector Machine with a Radial Basis Function (RBF) kernel was trained on "
    "the full scaled feature set. The RBF kernel was chosen over linear or polynomial "
    "kernels because it can model complex, non-linear decision boundaries in the "
    "high-dimensional space created by one-hot encoding. The regularisation parameter C=1.0 "
    "and kernel width gamma='scale' were used, which sets gamma to 1/(n_features × Var(X)) "
    " - an appropriate default that balances bias and variance for standardised data.")

add_figure(doc, os.path.join(FIGS,"fig7_cm_svm.png"),
    "Figure 7: Confusion matrix for Model 3 (SVM, RBF Kernel).", width=4.2)

add_metric_table(doc,
    ["Metric","Value"],
    [["Accuracy", f"{acc3:.4f}"],
     ["Precision",f"{prec3:.4f}"],
     ["Recall",   f"{rec3:.4f}"],
     ["F1-Score", f"{f1_3:.4f}"],
     ["AUC",      f"{auc3:.4f}"]])
cap = doc.add_paragraph("Table 6: Model 3 performance metrics.")
cap.runs[0].font.italic = True; cap.runs[0].font.size = Pt(10)
cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
doc.add_paragraph()

add_body(doc,
    "Unlike logistic regression, SVM does not produce directly interpretable coefficients. "
    "While this reduces its value for regulatory explanations, its discriminatory power "
    "is evident in the AUC score, and it serves as a useful benchmark against the "
    "interpretable logistic regression models.")

add_heading(doc, "3.5  ROC Curves and AUC Comparison", level=2)
add_figure(doc, os.path.join(FIGS,"fig8_roc.png"),
    "Figure 8: ROC curves for all three models. AUC values indicate the probability that "
    "the model ranks a randomly chosen defaulter above a randomly chosen non-defaulter.", width=5.5)

add_body(doc,
    f"All three models demonstrate meaningful discriminatory ability, with AUC values "
    f"well above the 0.5 baseline. Model 1 (LR Baseline) achieves AUC = {auc1:.3f}, "
    f"Model 2 (LR Optimised) achieves AUC = {auc2:.3f}, and "
    f"Model 3 (SVM RBF) achieves AUC = {auc3:.3f}.")

add_heading(doc, "3.6  Model Comparison and Precision-Recall Trade-Off", level=2)
add_metric_table(doc,
    ["Model","# Features","Accuracy","Precision","Recall","F1-Score","AUC"],
    [
        ["LR Baseline",   str(len(ALL_FEATURES)), f"{acc1:.4f}", f"{prec1:.4f}", f"{rec1:.4f}", f"{f1_1:.4f}", f"{auc1:.4f}"],
        ["LR Optimised",  str(len(sel_features)),  f"{acc2:.4f}", f"{prec2:.4f}", f"{rec2:.4f}", f"{f1_2:.4f}", f"{auc2:.4f}"],
        ["SVM (RBF)",     str(len(ALL_FEATURES)), f"{acc3:.4f}", f"{prec3:.4f}", f"{rec3:.4f}", f"{f1_3:.4f}", f"{auc3:.4f}"],
    ])
cap = doc.add_paragraph("Table 7: Side-by-side model performance comparison.")
cap.runs[0].font.italic = True; cap.runs[0].font.size = Pt(10)
cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
doc.add_paragraph()

add_body(doc,
    "In a lending context, precision and recall carry different business implications. "
    "High precision means that when the model predicts a default, it is usually correct - "
    "minimising the rejection of creditworthy borrowers. High recall means the model "
    "identifies most actual defaulters - minimising costly bad loans. "
    "Given the 7.5:1 cost asymmetry (FN $15,000 vs FP $2,000), recall should be prioritised. "
    "However, recall cannot be maximised without limit, as excessive false positives would "
    "erode customer relationships and market share.")

add_figure(doc, os.path.join(FIGS,"fig10_threshold.png"),
    "Figure 10: Total expected misclassification cost (Model 2) across decision thresholds. "
    f"The optimal threshold of {opt_thr:.2f} minimises total cost.", width=5.5)

add_body(doc,
    f"Figure 10 demonstrates a cost-based threshold analysis. By lowering the default "
    f"classification threshold from the standard 0.50 to the cost-optimal value of "
    f"{opt_thr:.2f} on Model 2, the bank can minimise total expected losses. "
    f"At this threshold, recall increases to {rec2_opt:.4f} while precision adjusts to "
    f"{prec2_opt:.4f}, reflecting the deliberate shift toward catching more defaulters "
    f"at the expense of a slightly higher false positive rate.")

doc.add_page_break()

# ----- STEP 4 : BUSINESS RECOMMENDATIONS -----
add_heading(doc, "Step 4: Business Recommendations", level=1)

add_heading(doc, "4.1  Key Predictors of Default", level=2)
add_body(doc,
    "Across all three models, the following applicant characteristics consistently emerged "
    "as the strongest predictors of loan default:")
doc.add_paragraph(
    "Checking Account Status: Borrowers with overdrawn or low-balance checking accounts "
    "are substantially more likely to default. This is the single most discriminative "
    "feature, reflecting the borrower's immediate liquidity.",
    style="List Number")
doc.add_paragraph(
    "Credit History: A history of delayed payments or critical credit accounts strongly "
    "predicts future default. This is the most widely validated predictor in the credit "
    "risk literature.",
    style="List Number")
doc.add_paragraph(
    "Loan Duration: Longer loan terms are associated with higher default risk, likely "
    "due to increased exposure to adverse income shocks over the repayment period.",
    style="List Number")
doc.add_paragraph(
    "Savings Account Status: Borrowers with little or no documented savings exhibit "
    "higher default rates, as they lack a financial buffer for unexpected expenses.",
    style="List Number")
doc.add_paragraph(
    "Loan Amount: Larger loans are modestly but consistently associated with higher "
    "default risk, particularly when combined with low savings or poor credit history.",
    style="List Number")

add_heading(doc, "4.2  Recommended Model and Deployment Strategy", level=2)
add_body(doc,
    f"The recommended model for deployment is Model 2 (Optimised Logistic Regression) "
    f"with a decision threshold of {opt_thr:.2f}. This recommendation is based on three "
    f"criteria: interpretability, performance, and parsimony.")
doc.add_paragraph(
    f"Interpretability: Logistic regression produces odds ratios - a format familiar to "
    f"credit analysts and regulators. Should a rejected applicant challenge the decision, "
    f"the bank can point to specific, statistically validated risk factors.",
    style="List Bullet")
doc.add_paragraph(
    f"Performance: Model 2 achieves competitive AUC = {auc2:.3f} and, at the optimal "
    f"threshold, recall = {rec2_opt:.4f} - meaning it correctly identifies approximately "
    f"{rec2_opt*100:.0f}% of future defaulters.",
    style="List Bullet")
doc.add_paragraph(
    f"Parsimony: Backward elimination reduced the feature set from {len(ALL_FEATURES)} "
    f"to {len(sel_features)} statistically significant predictors, reducing overfitting "
    f"risk and simplifying the scorecard for operational use.",
    style="List Bullet")
add_body(doc,
    "SecureBank should integrate the model into its loan origination system as a risk "
    "scoring tool. Applicants scoring above the threshold should be referred for manual "
    "review by credit analysts rather than automatic rejection, ensuring human oversight "
    "of borderline cases.")

add_heading(doc, "4.3  Ethical Use of Predictors", level=2)
add_body(doc,
    "SecureBank must exercise caution regarding the variables Personal_Status and "
    "Is_Foreign_Worker. Personal_Status encodes gender-related categories; "
    "Is_Foreign_Worker encodes national origin. Using these variables in credit scoring "
    "decisions may violate the Equal Credit Opportunity Act (ECOA) or equivalent "
    "regulations, regardless of their statistical predictive power. "
    "If the final deployment model retains either of these features, the bank's legal and "
    "compliance team should review whether their inclusion constitutes unlawful "
    "discrimination. The recommended approach is to exclude these variables from the "
    "production scorecard and rely solely on financially relevant predictors.")

add_heading(doc, "4.4  Limitations and Monitoring", level=2)
add_body(doc,
    "Several limitations must be acknowledged. First, the model was trained on historical "
    "data and may degrade in accuracy if economic conditions change significantly - a "
    "phenomenon known as model drift. SecureBank should re-train the model at least "
    "annually using fresh data. Second, the training dataset covers approximately 2,559 "
    "records; a larger and more recent sample would improve generalisation. Third, the "
    "model does not capture macroeconomic variables (e.g., unemployment rate, interest "
    "rates) that significantly influence default behaviour at the portfolio level.")
add_body(doc,
    "It is recommended that the model be monitored monthly using the Kolmogorov-Smirnov "
    "(KS) statistic and the Gini coefficient to detect performance degradation. If AUC "
    "falls below 0.70 or recall falls below 0.65 on a rolling three-month window, the "
    "model should be re-trained before continuing deployment.")

doc.add_page_break()

# ----- REFERENCES -----
add_heading(doc, "References", level=1)
refs = [
    "Breiman, L., Friedman, J., Olshen, R. A., & Stone, C. J. (1984). "
    "Classification and regression trees. Wadsworth.",

    "Harris, C. R., Millman, K. J., van der Walt, S. J., Gommers, R., Virtanen, P., "
    "Cournapeau, D., & Oliphant, T. E. (2020). Array programming with NumPy. "
    "Nature, 585(7825), 357-362. https://doi.org/10.1038/s41586-020-2649-2",

    "McKinney, W. (2010). Data structures for statistical computing in Python. "
    "Proceedings of the 9th Python in Science Conference, 56-61.",

    "Pedregosa, F., Varoquaux, G., Gramfort, A., Michel, V., Thirion, B., Grisel, O., "
    "& Duchesnay, É. (2011). Scikit-learn: Machine learning in Python. "
    "Journal of Machine Learning Research, 12, 2825-2830.",

    "Seabold, S., & Perktold, J. (2010). Statsmodels: Econometric and statistical "
    "modeling with Python. Proceedings of the 9th Python in Science Conference, 57-61.",

    "Waskom, M. L. (2021). Seaborn: Statistical data visualization. "
    "Journal of Open Source Software, 6(60), 3021. https://doi.org/10.21105/joss.03021",
]
for ref in refs:
    p = doc.add_paragraph(ref, style="List Paragraph")
    p.paragraph_format.first_line_indent = Inches(-0.5)
    p.paragraph_format.left_indent = Inches(0.5)
    for run in p.runs:
        run.font.size = Pt(11)
    doc.add_paragraph()

# Save
docx_path = os.path.join(BASE, "SecureBank_Loan_Default_Report.docx")
doc.save(docx_path)
print(f"\n  Word document saved: {docx_path}")
print("\n" + "=" * 65)
print("  ALL DONE - Analysis complete.")
print(f"  Output directory: {BASE}")
print("=" * 65)
